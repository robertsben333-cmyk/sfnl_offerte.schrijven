"""Tests for Word budget table section component."""
import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document

BASE = "skills/pptx_offerte/assets/sfnl_base.docx"
CONTENT = {
    "rows": [
        {"fase": "Fase 1", "dagen": 8, "kosten": 11840},
        {"fase": "Totaal", "dagen": 8, "kosten": 11840},
    ],
    "tarief_note": "Tarief: €1.480/dag",
    "termijnen": ["50% bij opdrachtverlening", "50% bij oplevering"],
}


def _full_text(doc) -> str:
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join([*(p.text for p in doc.paragraphs), *table_text])


def test_budget_table_word_contains_rows_and_terms():
    from skills.pptx_offerte.scripts.word.budget_table import add_section
    doc = Document(BASE)
    add_section(doc, CONTENT)
    text = _full_text(doc)
    assert "Fase 1" in text
    assert "Tarief: €1.480/dag" in text
    assert "opdrachtverlening" in text
