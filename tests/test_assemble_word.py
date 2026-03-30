"""Tests for assemble_word.py pipeline."""
import os, sys, pytest
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

BASE_DOCX = "skills/pptx_offerte/assets/sfnl_base.docx"


def test_assemble_word_empty_plan(tmp_path):
    """Assembling an empty plan produces a valid docx."""
    from skills.pptx_offerte.scripts.assemble_word import assemble_word
    from docx import Document

    output = str(tmp_path / "out.docx")
    assemble_word([], output, base=BASE_DOCX)
    doc = Document(output)
    assert doc is not None


def test_assemble_word_unknown_type_raises(tmp_path):
    """Unknown section type raises ValueError."""
    from skills.pptx_offerte.scripts.assemble_word import assemble_word

    output = str(tmp_path / "out.docx")
    with pytest.raises(ValueError, match="Unknown section type"):
        assemble_word([{"type": "nonexistent", "content": {}}], output, base=BASE_DOCX)


def test_assemble_word_accepts_plan_b_section_types(tmp_path):
    """Assembler accepts the new Plan B Word section components."""
    from skills.pptx_offerte.scripts.assemble_word import assemble_word
    from docx import Document

    output = str(tmp_path / "plan_b_out.docx")
    slide_plan = [
        {"type": "aanpak_section", "content": {
            "title": "Plan van aanpak",
            "subtitle": "Compacte samenvatting.",
            "phases": [{
                "naam": "Fase 1",
                "doel": "Doeltekst",
                "aanpak": "Aanpaktekst",
                "acties_sfnl": ["Actie 1"],
                "acties_klant": ["Actie 2"],
                "deliverable": "Rapport",
                "dagen": 8,
                "tijdlijn": "jan-feb 2026",
                "klant": "Testklant",
            }],
        }},
        {"type": "budget_table", "content": {
            "rows": [{"fase": "Fase 1", "dagen": 8, "kosten": 11840}],
            "termijnen": ["50% bij start"],
        }},
        {"type": "akkoord", "content": {
            "randvoorwaarden_tekst": "De teamsamenstelling is indicatief.",
            "termijnen": ["50% bij start"],
            "sfnl_naam": "Ruben Koekoek",
            "klant_naam": "Jan de Vries",
            "klant_org": "Testorganisatie",
        }},
    ]

    assemble_word(slide_plan, output, base=BASE_DOCX)
    doc = Document(output)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Plan van aanpak" in full_text
    assert "Ruben Koekoek" in full_text


def test_assemble_word_accepts_pptx_native_plan_types(tmp_path):
    """Word assembler can normalize the PPTX-native slide plan into linear sections."""
    from skills.pptx_offerte.scripts.assemble_word import assemble_word
    from docx import Document

    output = str(tmp_path / "normalized_out.docx")
    slide_plan = [
        {"type": "section_header", "content": {"title": "PLAN VAN AANPAK"}},
        {"type": "aanpak_overview", "content": {
            "title": "ONZE AANPAK",
            "subtitle": "In twee fases werken we naar inzicht.",
            "phases": [
                {"naam": "Fase 1", "beschrijving": "Eerste beschrijving", "tijdlijn": "jan-feb"},
                {"naam": "Fase 2", "beschrijving": "Tweede beschrijving", "tijdlijn": "mrt-apr"},
            ],
        }},
        {"type": "fase_detail", "content": {
            "number": 1,
            "naam": "Fase 1",
            "doel": "Doel fase 1",
            "aanpak": "Aanpak fase 1",
            "acties_sfnl": ["Deskresearch"],
            "acties_klant": ["Data delen"],
            "deliverable": "Rapport",
            "dagen": 6,
            "tijdlijn": "jan-feb 2026",
            "klant": "Testklant",
        }},
        {"type": "tijdslijn", "content": {
            "intro": "Indicatieve planning.",
            "disclaimer": "Planning is indicatief.",
            "phases": [{"naam": "Fase 2", "periode": "mrt-apr 2026", "activiteiten": "Validatie"}],
        }},
        {"type": "randvoorwaarden", "content": {"items": ["Eerste randvoorwaarde."]}},
        {"type": "akkoord", "content": {
            "sfnl_naam": "Ruben Koekoek",
            "klant_naam": "Jan de Vries",
            "klant_org": "Testorganisatie",
        }},
    ]

    assemble_word(slide_plan, output, base=BASE_DOCX)
    doc = Document(output)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ONZE AANPAK" in full_text
    assert "Doel fase 1" in full_text
    assert "Planning is indicatief." in full_text
    assert "Eerste randvoorwaarde." in full_text
