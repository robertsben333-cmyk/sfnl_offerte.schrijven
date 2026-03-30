"""Word aanpak section component."""
from docx import Document
from docx.shared import Pt


def _add_label_value(doc: Document, label: str, value: str) -> None:
    if not value:
        return
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(10)


def _add_bullet_lines(doc: Document, title: str, items: list[str]) -> None:
    if not items:
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(10)
    for item in items:
        paragraph = doc.add_paragraph(f"• {item}")
        if paragraph.runs:
            paragraph.runs[0].font.size = Pt(10)


def add_section(doc: Document, content: dict) -> None:
    """
    Add compact Word approach section.

    content keys:
      title (str)
      subtitle (str, optional)
      phases (list[dict])
    """
    doc.add_heading(content.get("title", "Plan van aanpak"), level=2)

    subtitle = content.get("subtitle", "")
    if subtitle:
        intro = doc.add_paragraph(subtitle)
        if intro.runs:
            intro.runs[0].italic = True

    for index, phase in enumerate(content.get("phases", []), start=1):
        number = phase.get("number", index)
        name = phase.get("naam", "")
        doc.add_heading(f"{number}. {name}".strip(), level=3)
        beschrijving = phase.get("beschrijving", "")
        if beschrijving:
            doc.add_paragraph(beschrijving)
        _add_label_value(doc, "Doel", phase.get("doel", ""))
        aanpak = phase.get("aanpak", "") or phase.get("beschrijving", "")
        _add_label_value(doc, "Aanpak", aanpak)
        _add_bullet_lines(doc, "Acties Social Finance NL", phase.get("acties_sfnl", []))

        klant = phase.get("klant", "Klant")
        _add_bullet_lines(doc, f"Acties {klant}", phase.get("acties_klant", []))
        _add_label_value(doc, "Deliverable", phase.get("deliverable", ""))
        _add_label_value(doc, "Tijdlijn", phase.get("tijdlijn", ""))
        dagen = phase.get("dagen", "")
        if dagen != "":
            _add_label_value(doc, "Dagdelen", str(dagen))

    timeline_note = content.get("timeline_note", "")
    if timeline_note:
        doc.add_paragraph(timeline_note)
