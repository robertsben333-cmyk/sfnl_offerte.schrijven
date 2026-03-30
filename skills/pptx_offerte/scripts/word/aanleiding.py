"""Word aanleiding section component."""
import json, os
from docx import Document

STYLE_PATH = os.path.join(os.path.dirname(__file__), "../../../../data/style.json")
with open(STYLE_PATH) as f:
    STYLE = json.load(f)

BLOCKS = [
    ("vraagstuk", "Maatschappelijk vraagstuk"),
    ("uitdagingen", "Grootste uitdagingen"),
    ("behoefte", "Behoefte van de klant"),
]


def add_section(doc: Document, content: dict) -> None:
    """
    Add aanleiding section with three labelled blocks.

    content keys: vraagstuk, uitdagingen, behoefte, summary_line (optional)
    """
    doc.add_heading("Aanleiding", level=2)

    summary = content.get("summary_line", "")
    if summary:
        p = doc.add_paragraph(summary)
        if p.runs:
            p.runs[0].italic = True

    for key, label in BLOCKS:
        doc.add_heading(label, level=3)
        doc.add_paragraph(content.get(key, ""))
