"""Word budget table section component."""
from docx import Document
from docx.shared import Pt


def _format_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = text
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].font.size = Pt(10)
        paragraph.runs[0].bold = bold


def add_section(doc: Document, content: dict) -> None:
    """
    Add Word budget section with table and notes.

    content keys:
      title (str, optional)
      rows (list[dict])
      tarief_note (str, optional)
      termijnen (list[str], optional)
    """
    doc.add_heading(content.get("title", "Begroting"), level=2)

    rows = content.get("rows", [])
    if rows:
        table = doc.add_table(rows=len(rows) + 1, cols=3)
        _format_cell(table.rows[0].cells[0], "Fase", bold=True)
        _format_cell(table.rows[0].cells[1], "Dagdelen", bold=True)
        _format_cell(table.rows[0].cells[2], "Kosten", bold=True)

        for row_index, row in enumerate(rows, start=1):
            _format_cell(table.rows[row_index].cells[0], str(row.get("fase", "")))
            _format_cell(table.rows[row_index].cells[1], str(row.get("dagen", "")))
            kosten = row.get("kosten", "")
            kosten_text = (
                f"€ {kosten:,.0f}".replace(",", ".")
                if isinstance(kosten, (int, float))
                else str(kosten)
            )
            _format_cell(table.rows[row_index].cells[2], kosten_text)

    tarief_note = content.get("tarief_note", "")
    if tarief_note:
        paragraph = doc.add_paragraph(tarief_note)
        if paragraph.runs:
            paragraph.runs[0].italic = True

    for termijn in content.get("termijnen", []):
        doc.add_paragraph(f"• {termijn}")
