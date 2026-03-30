"""Word akkoord section component."""
from docx import Document
from docx.shared import Pt


def add_section(doc: Document, content: dict) -> None:
    """
    Add Word akkoord section with terms and signature block.

    content keys:
      title (str, optional)
      randvoorwaarden_tekst (str, optional)
      termijnen (list[str], optional)
      sfnl_naam (str, optional)
      klant_naam (str, optional)
      klant_org (str, optional)
    """
    doc.add_heading(content.get("title", "Randvoorwaarden en akkoord"), level=2)

    randvoorwaarden = content.get("randvoorwaarden_tekst", "")
    if randvoorwaarden:
        doc.add_paragraph(randvoorwaarden)

    for item in content.get("randvoorwaarden_items", []):
        doc.add_paragraph(f"• {item}")

    termijnen = content.get("termijnen", [])
    if termijnen:
        heading = doc.add_paragraph()
        run = heading.add_run("Betaaltermijnen")
        run.bold = True
        run.font.size = Pt(10)
        for termijn in termijnen:
            doc.add_paragraph(f"• {termijn}")

    doc.add_paragraph("Voor akkoord:")
    sfnl_name = content.get("sfnl_naam", "")
    klant_name = content.get("klant_naam", "")
    klant_org = content.get("klant_org", "")

    left = doc.add_paragraph()
    left.add_run(sfnl_name).bold = True
    left.add_run(" — Social Finance NL")

    right = doc.add_paragraph()
    right.add_run(klant_name).bold = True
    if klant_org:
        right.add_run(f" — {klant_org}")
