"""Word cover section component."""
import json, os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

STYLE_PATH = os.path.join(os.path.dirname(__file__), "../../../../data/style.json")
with open(STYLE_PATH) as f:
    STYLE = json.load(f)


def add_section(doc: Document, content: dict) -> None:
    """
    Add a cover page section.

    content keys: client, title, date, proposition
    """
    doc.add_paragraph("")  # spacer

    p_client = doc.add_paragraph(content.get("client", ""))
    if p_client.runs:
        p_client.runs[0].font.size = Pt(14)
        p_client.runs[0].font.name = STYLE["fonts"]["body"]

    p_title = doc.add_heading(content.get("title", ""), level=1)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_label = doc.add_paragraph("OFFERTE")
    if p_label.runs:
        p_label.runs[0].font.size = Pt(10)
        p_label.runs[0].bold = True

    p_date = doc.add_paragraph(content.get("date", ""))
    if p_date.runs:
        p_date.runs[0].font.size = Pt(11)

    doc.add_page_break()
