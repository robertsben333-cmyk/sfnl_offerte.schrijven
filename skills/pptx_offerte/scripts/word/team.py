"""Word team section component."""
import json, os
from docx import Document
from docx.shared import Pt

STYLE_PATH = os.path.join(os.path.dirname(__file__), "../../../../data/style.json")
with open(STYLE_PATH) as f:
    STYLE = json.load(f)


def add_section(doc: Document, content: dict) -> None:
    """
    Add team section.

    content keys:
      members (list): Each has name, title, bio
    """
    doc.add_heading("Ons team", level=2)

    for member in content.get("members", []):
        p_name = doc.add_paragraph()
        run = p_name.add_run(member.get("name", ""))
        run.bold = True
        run.font.size = Pt(12)

        p_title = doc.add_paragraph(member.get("title", ""))
        if p_title.runs:
            p_title.runs[0].font.size = Pt(10)
            p_title.runs[0].italic = True

        doc.add_paragraph(member.get("bio", ""))
        doc.add_paragraph("")  # spacer
